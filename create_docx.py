from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='宋体', font_size=10.5, bold=False):
    """设置中文字体"""
    font = run.font
    font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    font.size = Pt(font_size)
    font.bold = bold

def add_heading_zh(doc, text, level=1):
    """添加中文标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    font_names = ['黑体', 'SimHei'] if level <= 2 else ['宋体', 'SimSun']
    set_chinese_font(run, font_name=font_names[0], font_size=(18 if level==1 else (14 if level==2 else 12)), bold=True)
    return heading

def add_paragraph_zh(doc, text, bold=False, font_size=10.5):
    """添加中文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, font_name='宋体', font_size=font_size, bold=bold)
    return p

def add_code_block(doc, code_text):
    """添加代码块"""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    set_chinese_font(run, font_name='Consolas', font_size=9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def main():
    doc = Document()
    
    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)
    
    # 标题
    title = doc.add_heading('AI调度中心 - 企业API Key管理系统', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('\n启动说明文档')
    set_chinese_font(run, font_name='黑体', font_size=22, bold=True)
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('（一键启动脚本使用指南）')
    set_chinese_font(run, font_name='宋体', font_size=12)
    
    doc.add_paragraph()
    
    # 系统概述
    add_heading_zh(doc, '一、系统概述', level=1)
    add_paragraph_zh(doc, 'AI调度中心是一个开源企业级大模型API Key管理系统，采用前后端分离架构：')
    
    # 架构说明表格
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    headers = ['组件', '技术栈']
    data = [
        ['前端', 'Vue 3 + Vite + Element Plus'],
        ['后端', 'Spring Boot + JPA + MySQL + Redis'],
        ['一键启动', 'start.ps1（PowerShell脚本）']
    ]
    
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(header)
        set_chinese_font(run, font_name='黑体', font_size=10.5, bold=True)
    
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i+1].cells[j]
            run = cell.paragraphs[0].add_run(cell_text)
            set_chinese_font(run, font_name='宋体', font_size=10.5)
    
    doc.add_paragraph()
    
    # 准备工作
    add_heading_zh(doc, '二、准备工作', level=1)
    
    add_heading_zh(doc, '2.1 解压项目文件', level=2)
    add_paragraph_zh(doc, '将收到的压缩包解压到本地目录，确保目录结构如下：')
    
    code = '''AI调度中心——企业API Key管理系统/
├── start.ps1              # 一键启动脚本（核心文件）
├── init-db.bat            # 数据库初始化脚本
├── docker-compose.yml     # Docker编排配置
├── Dockerfile             # 后端Docker镜像构建文件
├── backend/               # 后端项目目录
│   ├── pom.xml
│   └── src/
└── frontend/              # 前端项目目录
    ├── package.json
    └── src/'''
    add_code_block(doc, code)
    
    add_heading_zh(doc, '2.2 系统环境要求', level=2)
    add_paragraph_zh(doc, '以下环境满足其一即可启动系统：')
    
    # 环境要求表格
    table2 = doc.add_table(rows=9, cols=3)
    table2.style = 'Light Grid Accent 1'
    
    headers2 = ['组件', '版本要求', '用途']
    data2 = [
        ['Windows', 'Windows 10/11', '操作系统'],
        ['PowerShell', '5.1+', '运行启动脚本'],
        ['Node.js', '18+', '前端运行环境'],
        ['Java', '17+', '后端运行环境'],
        ['Maven', '3.8+', '后端构建工具'],
        ['Docker Desktop', '最新版', '容器化运行（推荐）'],
        ['MySQL', '8.0', '数据库（可用Docker替代）'],
        ['Redis', '7.x', '缓存（可用Docker替代）']
    ]
    
    for i, header in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(header)
        set_chinese_font(run, font_name='黑体', font_size=10.5, bold=True)
    
    for i, row_data in enumerate(data2):
        for j, cell_text in enumerate(row_data):
            cell = table2.rows[i+1].cells[j]
            run = cell.paragraphs[0].add_run(cell_text)
            set_chinese_font(run, font_name='宋体', font_size=10.5)
    
    doc.add_paragraph()
    
    # 重要提示
    p = doc.add_paragraph()
    run = p.add_run('重要提示：')
    set_chinese_font(run, font_name='黑体', font_size=10.5, bold=True)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    run = p.add_run('至少需要安装 Node.js + Java + Maven 或 Docker Desktop 其中一种组合。如果只安装Docker Desktop，系统会自动用Docker容器运行所有服务，无需单独安装Java和Maven。')
    set_chinese_font(run, font_name='宋体', font_size=10.5)
    
    doc.add_page_break()
    
    # 环境安装指南
    add_heading_zh(doc, '三、环境安装指南', level=1)
    
    add_heading_zh(doc, '3.1 安装 Node.js（前端必需）', level=2)
    steps = [
        '访问 https://nodejs.org/',
        '下载 LTS（长期支持）版本（推荐 v20.x）',
        '运行安装程序，一路点击"Next"完成安装',
        '验证安装：打开PowerShell，执行以下命令：'
    ]
    for step in steps:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(step)
        set_chinese_font(run, font_name='宋体', font_size=10.5)
    
    add_code_block(doc, 'node -v    # 应显示版本号，如 v20.11.0\nnpm -v     # 应显示版本号，如 10.2.4')
    
    add_heading_zh(doc, '3.2 安装 Java（后端必需，如无Docker）', level=2)
    steps2 = [
        '访问 https://adoptium.net/',
        '下载 OpenJDK 17 LTS（Windows x64 MSI Installer）',
        '运行安装程序，勾选"Add to PATH"',
        '验证安装：'
    ]
    for step in steps2:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(step)
        set_chinese_font(run, font_name='宋体', font_size=10.5)
    
    add_code_block(doc, 'java -version    # 应显示 openjdk version "17.x.x"')
    
    add_heading_zh(doc, '3.3 安装 Maven（后端必需，如无Docker）', level=2)
    steps3 = [
        '访问 https://maven.apache.org/download.cgi',
        '下载 Binary zip archive（如 apache-maven-3.9.x-bin.zip）',
        '解压到 C:\\apache-maven-3.9.x',
        '配置环境变量：右键"此电脑" → 属性 → 高级系统设置 → 环境变量',
        '在"系统变量"中找到 Path，点击编辑，添加：C:\\apache-maven-3.9.x\\bin',
        '验证安装：'
    ]
    for step in steps3:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(step)
        set_chinese_font(run, font_name='宋体', font_size=10.5)
    
    add_code_block(doc, 'mvn -v    # 应显示 Apache Maven 3.9.x')
    
    add_heading_zh(doc, '3.4 安装 Docker Desktop（推荐，最简单）', level=2)
    steps4 = [
        '访问 https://www.docker.com/products/docker-desktop/',
        '下载并安装 Docker Desktop',
        '启动 Docker Desktop，等待显示"Docker Desktop is running"',
        '验证安装：'
    ]
    for step in steps4:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(step)
        set_chinese_font(run, font_name='宋体', font_size=10.5)
    
    add_code_block(doc, 'docker --version\ndocker-compose --version')
    
    p = doc.add_paragraph()
    run = p.add_run('推荐方案：如果只安装 Docker Desktop，系统会自动用Docker容器运行MySQL、Redis和后端，无需单独安装Java和Maven。这是最简单的部署方式。')
    set_chinese_font(run, font_name='宋体', font_size=10.5)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)
    
    doc.add_page_break()
    
    # 启动系统
    add_heading_zh(doc, '四、启动系统', level=1)
    
    add_heading_zh(doc, '4.1 使用一键启动脚本（推荐）', level=2)
    
    add_paragraph_zh(doc, '步骤1：打开 PowerShell', bold=True)
    add_paragraph_zh(doc, '按 Win + X，选择 Windows PowerShell 或 终端。建议右键点击开始菜单中的 PowerShell，选择"以管理员身份运行"。')
    
    add_paragraph_zh(doc, '步骤2：进入项目目录', bold=True)
    add_code_block(doc, 'cd "C:\\Users\\你的用户名\\Desktop\\AI调度中心——企业API Key管理系统"')
    add_paragraph_zh(doc, '（根据实际解压路径修改）')
    
    add_paragraph_zh(doc, '步骤3：执行启动脚本', bold=True)
    add_code_block(doc, '.\\start.ps1')
    
    p = doc.add_paragraph()
    run = p.add_run('注意：')
    set_chinese_font(run, font_name='黑体', font_size=10.5, bold=True)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    run = p.add_run('如果提示"执行策略禁止运行脚本"，请先执行以下命令：')
    set_chinese_font(run, font_name='宋体', font_size=10.5)
    
    add_code_block(doc, 'Set-ExecutionPolicy -ExecutionPolicy RemoteSigned -Scope CurrentUser')
    
    add_paragraph_zh(doc, '步骤4：等待启动完成', bold=True)
    add_paragraph_zh(doc, '脚本会自动执行以下操作：')
    
    auto_steps = [
        '检测环境（Node.js、Java、Maven、Docker）',
        '启动 MySQL（本地或Docker）',
        '启动 Redis（本地或Docker）',
        '初始化数据库',
        '检查端口占用情况',
        '编译并启动后端服务',
        '安装依赖并启动前端服务',
        '自动打开浏览器'
    ]
    for step in auto_steps:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(step)
        set_chinese_font(run, font_name='宋体', font_size=10.5)
    
    add_heading_zh(doc, '4.2 启动过程说明', level=2)
    add_paragraph_zh(doc, '脚本执行时会显示以下步骤：')
    
    add_code_block(doc, '''[1/9] Detecting environment...          # 检测环境
[2/9] Checking infrastructure...        # 检查MySQL
[3/9] Checking infrastructure...        # 检查Redis
[4/9] Initializing database...          # 初始化数据库
[5/9] Checking port availability...     # 检查端口占用
[6/9] Starting backend service...       # 启动后端
[7/9] Starting frontend service...      # 启动前端
[8/9] Opening browser...                # 打开浏览器
[9/9] System startup complete!          # 启动完成''')
    
    add_heading_zh(doc, '4.3 访问系统', level=2)
    add_paragraph_zh(doc, '启动成功后，浏览器会自动打开前端页面。您也可以通过以下地址访问：')
    
    # 访问地址表格
    table3 = doc.add_table(rows=5, cols=2)
    table3.style = 'Light Grid Accent 1'
    
    headers3 = ['服务', '访问地址']
    data3 = [
        ['前端页面', 'http://localhost:5173'],
        ['后端API', 'http://localhost:8080'],
        ['接口文档(Swagger)', 'http://localhost:8080/swagger-ui/index.html'],
        ['健康检查', 'http://localhost:8080/api/health']
    ]
    
    for i, header in enumerate(headers3):
        cell = table3.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(header)
        set_chinese_font(run, font_name='黑体', font_size=10.5, bold=True)
    
    for i, row_data in enumerate(data3):
        for j, cell_text in enumerate(row_data):
            cell = table3.rows[i+1].cells[j]
            run = cell.paragraphs[0].add_run(cell_text)
            set_chinese_font(run, font_name='宋体', font_size=10.5)
    
    doc.add_paragraph()
    
    add_paragraph_zh(doc, '默认登录账号：', bold=True)
    add_paragraph_zh(doc, '用户名：admin')
    add_paragraph_zh(doc, '密码：admin123')
    
    doc.add_page_break()
    
    # 常见问题解决
    add_heading_zh(doc, '五、常见问题解决', level=1)
    
    issues = [
        ('端口被占用', 
         '脚本提示 "Port 8080 occupied" 或 "Port 5173 occupied"',
         '1. 脚本会询问是否自动终止占用进程\n2. 输入 Y 让脚本自动处理\n3. 或输入 N 后手动关闭占用端口的程序'),
        ('MySQL连接失败', 
         '提示 "MySQL not running"',
         '1. 如果已安装 Docker Desktop，脚本会自动启动MySQL容器\n2. 如果没有Docker，请先安装并启动本地MySQL\n3. 确保MySQL root密码为 root（可在 application.yml 中修改）'),
        ('前端依赖安装失败', 
         'npm install 卡住或报错',
         '1. 检查网络连接\n2. 尝试更换npm镜像源：npm config set registry https://registry.npmmirror.com\n3. 删除 frontend/node_modules 目录后重新运行脚本'),
        ('后端编译失败', 
         'Maven编译报错',
         '1. 检查Java和Maven是否正确安装\n2. 检查 backend/pom.xml 是否存在\n3. 尝试手动编译：cd backend && mvn clean compile'),
        ('Docker启动失败', 
         'Docker容器无法启动',
         '1. 确保 Docker Desktop 正在运行\n2. 检查 Docker 设置中的资源分配（内存建议4G+）\n3. 尝试重启 Docker Desktop')
    ]
    
    for i, (title, symptom, solution) in enumerate(issues, 1):
        add_heading_zh(doc, f'5.{i} {title}', level=2)
        
        p = doc.add_paragraph()
        run = p.add_run('现象：')
        set_chinese_font(run, font_name='黑体', font_size=10.5, bold=True)
        run = p.add_run(symptom)
        set_chinese_font(run, font_name='宋体', font_size=10.5)
        
        p = doc.add_paragraph()
        run = p.add_run('解决：')
        set_chinese_font(run, font_name='黑体', font_size=10.5, bold=True)
        
        for line in solution.split('\n'):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(line)
            set_chinese_font(run, font_name='宋体', font_size=10.5)
    
    doc.add_page_break()
    
    # 手动启动方式
    add_heading_zh(doc, '六、手动启动方式（备用）', level=1)
    add_paragraph_zh(doc, '如果一键脚本无法正常工作，可以手动分步启动：')
    
    add_heading_zh(doc, '6.1 方式一：Docker Compose 启动（最简单）', level=2)
    add_code_block(doc, '''# 进入项目目录
cd "项目根目录"

# 启动所有服务（MySQL + Redis + 后端）
docker-compose up -d

# 等待1-2分钟后，手动启动前端
cd frontend
npm install
npm run dev''')
    
    add_heading_zh(doc, '6.2 方式二：完全手动启动', level=2)
    
    add_paragraph_zh(doc, '步骤1：启动MySQL', bold=True)
    add_paragraph_zh(doc, '使用本地MySQL：启动MySQL服务，创建数据库 ai_key_management')
    add_paragraph_zh(doc, '或使用Docker：')
    add_code_block(doc, 'docker run -d --name aikey-mysql -p 3306:3306 -e MYSQL_ROOT_PASSWORD=root -e MYSQL_DATABASE=ai_key_management mysql:8.0')
    
    add_paragraph_zh(doc, '步骤2：启动Redis（可选）', bold=True)
    add_code_block(doc, 'docker run -d --name aikey-redis -p 6379:6379 redis:7-alpine')
    
    add_paragraph_zh(doc, '步骤3：初始化数据库', bold=True)
    add_code_block(doc, '# 使用提供的脚本\\n.\\init-db.bat\\n# 选择选项1进行完整初始化')
    
    add_paragraph_zh(doc, '步骤4：启动后端', bold=True)
    add_code_block(doc, 'cd backend\nmvn clean compile\nmvn spring-boot:run')
    
    add_paragraph_zh(doc, '步骤5：启动前端', bold=True)
    add_code_block(doc, 'cd frontend\nnpm install\nnpm run dev')
    
    doc.add_page_break()
    
    # 系统功能简介
    add_heading_zh(doc, '七、系统功能简介', level=1)
    add_paragraph_zh(doc, '启动成功后，您可以使用以下功能：')
    
    # 功能表格
    table4 = doc.add_table(rows=9, cols=2)
    table4.style = 'Light Grid Accent 1'
    
    headers4 = ['功能模块', '说明']
    data4 = [
        ['渠道管理', '管理大模型API渠道（OpenAI、Claude等）'],
        ['模型市场', '配置和管理可用模型'],
        ['项目管理', '创建和管理API Key项目'],
        ['虚拟Key管理', '生成和管理虚拟API Key'],
        ['真实Key管理', '管理上游真实API Key'],
        ['额度管理', '分配和监控API调用额度'],
        ['调用日志', '查看API调用记录和统计'],
        ['团队管理', '多团队协作和权限管理']
    ]
    
    for i, header in enumerate(headers4):
        cell = table4.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(header)
        set_chinese_font(run, font_name='黑体', font_size=10.5, bold=True)
    
    for i, row_data in enumerate(data4):
        for j, cell_text in enumerate(row_data):
            cell = table4.rows[i+1].cells[j]
            run = cell.paragraphs[0].add_run(cell_text)
            set_chinese_font(run, font_name='宋体', font_size=10.5)
    
    doc.add_page_break()
    
    # 停止系统
    add_heading_zh(doc, '八、停止系统', level=1)
    
    add_heading_zh(doc, '8.1 停止一键启动脚本', level=2)
    add_paragraph_zh(doc, '在 PowerShell 窗口中按 Ctrl + C，脚本会自动停止所有服务。')
    
    add_heading_zh(doc, '8.2 手动停止', level=2)
    add_paragraph_zh(doc, '停止后端：在运行后端的终端按 Ctrl+C')
    add_paragraph_zh(doc, '停止前端：在运行前端的终端按 Ctrl+C')
    add_paragraph_zh(doc, '停止Docker容器：')
    add_code_block(doc, 'docker-compose down')
    
    # 技术支持
    add_heading_zh(doc, '九、技术支持', level=1)
    add_paragraph_zh(doc, '如遇到问题，请检查以下清单：')
    
    add_paragraph_zh(doc, '环境检查清单：', bold=True)
    checklist = [
        'Node.js 已安装 (node -v)',
        'Java 已安装 (java -version)',
        'Maven 已安装 (mvn -v)',
        'Docker Desktop 正在运行（如使用Docker）'
    ]
    for item in checklist:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        set_chinese_font(run, font_name='宋体', font_size=10.5)
    
    add_paragraph_zh(doc, '端口检查：', bold=True)
    ports = [
        '3306 端口未被占用（MySQL）',
        '6379 端口未被占用（Redis）',
        '8080 端口未被占用（后端）',
        '5173 端口未被占用（前端）'
    ]
    for item in ports:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        set_chinese_font(run, font_name='宋体', font_size=10.5)
    
    add_paragraph_zh(doc, '日志查看：', bold=True)
    logs = [
        '后端日志：查看 PowerShell 输出',
        'Docker日志：docker logs aikey-backend',
        'MySQL日志：docker logs aikey-mysql'
    ]
    for item in logs:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        set_chinese_font(run, font_name='宋体', font_size=10.5)
    
    # 文件清单
    add_heading_zh(doc, '十、文件清单', level=1)
    add_paragraph_zh(doc, '确保您有以下完整文件：')
    
    add_code_block(doc, '''项目根目录/
├── start.ps1              ✓ 一键启动脚本
├── init-db.bat            ✓ 数据库初始化脚本
├── docker-compose.yml     ✓ Docker编排文件
├── Dockerfile             ✓ Docker镜像构建文件
├── backend/               ✓ 后端代码目录
│   ├── pom.xml
│   └── src/main/resources/db/
│       ├── schema.sql     ✓ 数据库表结构
│       └── data.sql       ✓ 初始数据
└── frontend/              ✓ 前端代码目录
    ├── package.json
    └── src/''')
    
    doc.add_paragraph()
    
    # 结束语
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('祝您使用愉快！')
    set_chinese_font(run, font_name='黑体', font_size=14, bold=True)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x99)
    
    # 保存文档
    output_path = r'C:\Users\26404\Desktop\26计设材料\AI调度中心系统启动说明.docx'
    doc.save(output_path)
    print(f'Word文档已生成：{output_path}')

if __name__ == '__main__':
    main()
