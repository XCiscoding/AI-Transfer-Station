from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 创建文档
doc = Document()

# 设置中文字体
def set_chinese_font(run, font_name='SimSun', size=12, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold

# 添加标题
title = doc.add_heading('', level=0)
title_run = title.add_run('AI调度中心 - 完整部署指南')
title_run.font.size = Pt(22)
title_run.font.bold = True
set_chinese_font(title_run, 'SimHei', 22, True)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加说明信息
info = doc.add_paragraph()
info_run = info.add_run('目标读者：未接触过此项目的人\n部署方式：单台云服务器 + Docker + FinalShell\n预计时间：30-45分钟\n难度等级：⭐⭐ 简单')
set_chinese_font(info_run, 'SimSun', 11)
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加分隔线
doc.add_paragraph('─' * 50)

# 目录
doc.add_heading('目录', level=1)
toc_items = [
    '一、准备工作',
    '二、购买云服务器',
    '三、配置安全组',
    '四、连接服务器',
    '五、上传项目文件',
    '六、执行部署',
    '七、验证部署',
    '八、常见问题',
    '九、日常维护',
    '十、获取帮助',
    '附录：文件清单'
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Number')
    for run in p.runs:
        set_chinese_font(run, 'SimSun', 11)

doc.add_page_break()

# 一、准备工作
doc.add_heading('一、准备工作', level=1)

doc.add_heading('1.1 你需要准备什么？', level=2)
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '物品'
hdr_cells[1].text = '用途'
hdr_cells[2].text = '获取方式'

items = [
    ['云服务器', '运行项目', '阿里云/腾讯云/华为云等'],
    ['FinalShell', '连接服务器', '官网下载(http://www.hostbuf.com/)'],
    ['本项目代码', '部署内容', '从开发者处获取']
]
for i, item in enumerate(items, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = item[0]
    row_cells[1].text = item[1]
    row_cells[2].text = item[2]

# 设置表格字体
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, 'SimSun', 10)

doc.add_heading('1.2 项目架构说明', level=2)

arch_text = '''┌─────────────────────────────────────────┐
│           云服务器 (2核4G)               │
│  ┌─────────────────────────────────┐    │
│  │         Docker 容器化运行         │    │
│  │  ┌─────────┐  ┌─────────────┐   │    │
│  │  │ 前端服务 │  │  后端服务    │   │    │
│  │  │ (Nginx) │  │(Spring Boot)│   │    │
│  │  │ :80端口 │  │  :8080端口   │   │    │
│  │  └────┬────┘  └──────┬──────┘   │    │
│  │       │              │          │    │
│  │  ┌────┴──────────────┴──────┐   │    │
│  │  │      MySQL 数据库         │   │    │
│  │  │   (ai_key_management)    │   │    │
│  │  └──────────────────────────┘   │    │
│  │  ┌──────────────────────────┐   │    │
│  │  │      Redis 缓存          │   │    │
│  │  └──────────────────────────┘   │    │
│  └─────────────────────────────────┘    │
└─────────────────────────────────────────┘'''

p = doc.add_paragraph()
run = p.add_run(arch_text)
set_chinese_font(run, 'Consolas', 9)

doc.add_paragraph('技术栈：', style='List Bullet')
tech_stack = [
    '前端：Vue 3 + Vite + Element Plus',
    '后端：Spring Boot 3.2 + Java 17',
    '数据库：MySQL 8.0',
    '缓存：Redis 7',
    '容器化：Docker + Docker Compose'
]
for item in tech_stack:
    p = doc.add_paragraph(item, style='List Bullet 2')
    for run in p.runs:
        set_chinese_font(run, 'SimSun', 10)

doc.add_page_break()

# 二、购买云服务器
doc.add_heading('二、购买云服务器', level=1)

doc.add_heading('2.1 推荐配置', level=2)
table = doc.add_table(rows=6, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '配置项'
hdr_cells[1].text = '推荐规格'
hdr_cells[2].text = '说明'

configs = [
    ['CPU', '2核', '足够运行所有服务'],
    ['内存', '4GB', '最低要求，建议4G以上'],
    ['带宽', '3-5M', '访问速度有保障'],
    ['系统盘', '40GB+', '存放代码和数据库'],
    ['操作系统', 'CentOS 7/8 或 Ubuntu 20.04+', '推荐 CentOS 7.9']
]
for i, config in enumerate(configs, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = config[0]
    row_cells[1].text = config[1]
    row_cells[2].text = config[2]

for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, 'SimSun', 10)

doc.add_heading('2.2 购买步骤（以阿里云为例）', level=2)
steps = [
    '访问阿里云官网 (https://www.aliyun.com/)',
    '搜索"云服务器ECS" → 点击"立即购买"',
    '选择配置：\n   - 计费方式：包年包月（长期使用）或按量付费（测试）\n   - 地域：选择离你最近的\n   - 实例规格：2核4G（共享标准型s6 或 突发性能实例t6）\n   - 镜像：CentOS 7.9 64位\n   - 带宽：3Mbps\n   - 安全组：先不设置，后面单独配置',
    '设置登录密码（记住这个密码！）',
    '确认订单并支付'
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph(f'{i}. {step}')
    for run in p.runs:
        set_chinese_font(run, 'SimSun', 10)

p = doc.add_paragraph()
run = p.add_run('💡 省钱提示：新用户首年通常有优惠，约 200-400元/年')
set_chinese_font(run, 'SimSun', 10)
run.font.color.rgb = RGBColor(0, 128, 0)

doc.add_page_break()

# 三、配置安全组
doc.add_heading('三、配置安全组', level=1)

doc.add_heading('3.1 什么是安全组？', level=2)
p = doc.add_paragraph('安全组 = 防火墙规则，决定哪些端口可以被外部访问')
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

doc.add_heading('3.2 需要开放的端口', level=2)
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '端口'
hdr_cells[1].text = '用途'
hdr_cells[2].text = '必须开放？'

ports = [
    ['22', 'SSH远程连接', '✅ 必须'],
    ['80', 'HTTP网页访问', '✅ 必须'],
    ['443', 'HTTPS（暂不需要）', '❌ 可选']
]
for i, port in enumerate(ports, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = port[0]
    row_cells[1].text = port[1]
    row_cells[2].text = port[2]

for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, 'SimSun', 10)

doc.add_heading('3.3 配置步骤（阿里云为例）', level=2)
steps = [
    '进入 ECS控制台 (https://ecs.console.aliyun.com/)',
    '找到你的服务器 → 点击"安全组"',
    '点击"配置规则"',
    '点击"入方向" → "手动添加"',
    '添加以下规则：'
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph(f'{i}. {step}')
    for run in p.runs:
        set_chinese_font(run, 'SimSun', 10)

p = doc.add_paragraph()
run = p.add_run('规则1：\n- 授权策略：允许\n- 优先级：1\n- 协议类型：自定义TCP\n- 端口范围：22/22\n- 授权对象：0.0.0.0/0\n- 描述：SSH远程连接\n\n规则2：\n- 授权策略：允许\n- 优先级：1\n- 协议类型：自定义TCP\n- 端口范围：80/80\n- 授权对象：0.0.0.0/0\n- 描述：HTTP网页访问')
set_chinese_font(run, 'Consolas', 9)

p = doc.add_paragraph()
run = p.add_run('⚠️ 注意：0.0.0.0/0 表示允许所有IP访问，生产环境建议限制为你的IP')
set_chinese_font(run, 'SimSun', 10)
run.font.color.rgb = RGBColor(255, 0, 0)

doc.add_page_break()

# 四、连接服务器
doc.add_heading('四、连接服务器', level=1)

doc.add_heading('4.1 安装 FinalShell', level=2)
steps = [
    '访问 FinalShell官网 (http://www.hostbuf.com/)',
    '下载对应系统的版本（Windows/Mac/Linux）',
    '安装并打开'
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph(f'{i}. {step}')
    for run in p.runs:
        set_chinese_font(run, 'SimSun', 10)

doc.add_heading('4.2 创建连接', level=2)
p = doc.add_paragraph('1. 点击左上角 "文件夹+" 图标 → "SSH连接"')
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

p = doc.add_paragraph('2. 填写连接信息：')
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

p = doc.add_paragraph()
run = p.add_run('名称：AI调度中心服务器\n主机：你的服务器公网IP（如 123.45.67.89）\n端口：22\n用户名：root\n密码：你购买时设置的密码')
set_chinese_font(run, 'Consolas', 9)

steps = [
    '点击"确定"',
    '双击连接名称，首次连接选择"接受并保存"'
]
for i, step in enumerate(steps, 3):
    p = doc.add_paragraph(f'{i}. {step}')
    for run in p.runs:
        set_chinese_font(run, 'SimSun', 10)

doc.add_heading('4.3 连接成功标志', level=2)
p = doc.add_paragraph('看到类似下面的提示表示成功：')
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

p = doc.add_paragraph()
run = p.add_run('[root@iZbp123456789 ~]#')
set_chinese_font(run, 'Consolas', 10)

doc.add_page_break()

# 五、上传项目文件
doc.add_heading('五、上传项目文件', level=1)

doc.add_heading('5.1 需要上传哪些文件？', level=2)
p = doc.add_paragraph('从本地项目文件夹中，上传以下内容到服务器：')
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

p = doc.add_paragraph()
run = p.add_run('''必须上传的文件：
├── backend/              # 后端代码文件夹
│   ├── src/
│   ├── pom.xml
│   └── ...
├── frontend/             # 前端代码文件夹
│   ├── src/
│   ├── package.json
│   └── ...
└── deploy/               # 部署配置文件夹
    ├── docker-compose.all-in-one.yml
    ├── Dockerfile.backend
    ├── Dockerfile.frontend
    ├── nginx.conf
    ├── finalshell-deploy.sh
    └── .env.example''')
set_chinese_font(run, 'Consolas', 9)

doc.add_heading('5.2 上传步骤', level=2)
p = doc.add_paragraph('方法：使用 FinalShell 文件管理器')
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10, True)

steps = [
    '在 FinalShell 左侧找到"文件管理器"',
    '本地：找到你的项目文件夹',
    '远程：进入 /opt 目录（命令：cd /opt）',
    '创建项目目录：mkdir aikey',
    '进入目录：cd aikey',
    '拖拽上传：\n   - 将 backend 文件夹拖到远程 /opt/aikey/\n   - 将 frontend 文件夹拖到远程 /opt/aikey/\n   - 将 deploy 文件夹拖到远程 /opt/aikey/'
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph(f'{i}. {step}')
    for run in p.runs:
        set_chinese_font(run, 'SimSun', 10)

doc.add_heading('5.3 验证上传', level=2)
p = doc.add_paragraph('在 FinalShell 终端执行：')
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

p = doc.add_paragraph()
run = p.add_run('cd /opt/aikey\nls -la')
set_chinese_font(run, 'Consolas', 10)

p = doc.add_paragraph('应该看到：')
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

p = doc.add_paragraph()
run = p.add_run('''drwxr-xr-x 5 root root 4096 Jan 15 10:00 .
drwxr-xr-x 3 root root 4096 Jan 15 09:59 ..
drwxr-xr-x 6 root root 4096 Jan 15 10:00 backend
drwxr-xr-x 6 root root 4096 Jan 15 10:00 deploy
drwxr-xr-x 6 root root 4096 Jan 15 10:00 frontend''')
set_chinese_font(run, 'Consolas', 9)

doc.add_page_break()

# 六、执行部署
doc.add_heading('六、执行部署', level=1)

doc.add_heading('6.1 进入部署目录', level=2)
p = doc.add_paragraph()
run = p.add_run('cd /opt/aikey/deploy')
set_chinese_font(run, 'Consolas', 10)

doc.add_heading('6.2 运行一键部署脚本', level=2)
p = doc.add_paragraph()
run = p.add_run('bash finalshell-deploy.sh')
set_chinese_font(run, 'Consolas', 10)

doc.add_heading('6.3 部署过程说明', level=2)
p = doc.add_paragraph('脚本会自动执行以下步骤：')
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

table = doc.add_table(rows=8, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '步骤'
hdr_cells[1].text = '耗时'
hdr_cells[2].text = '说明'

deploy_steps = [
    ['1. 检查环境', '5秒', '检查系统版本和权限'],
    ['2. 安装Docker', '2-5分钟', '自动安装Docker和Docker Compose'],
    ['3. 准备目录', '5秒', '创建必要的文件夹'],
    ['4. 配置环境变量', '5秒', '自动生成数据库密码'],
    ['5. 构建镜像', '10-20分钟', '编译前端和后端代码'],
    ['6. 启动服务', '1-2分钟', '启动所有容器'],
    ['7. 健康检查', '30秒', '验证服务是否正常']
]
for i, step in enumerate(deploy_steps, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = step[0]
    row_cells[1].text = step[1]
    row_cells[2].text = step[2]

for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, 'SimSun', 10)

p = doc.add_paragraph()
run = p.add_run('总耗时：约15-30分钟')
set_chinese_font(run, 'SimSun', 11, True)

doc.add_heading('6.4 部署成功标志', level=2)
p = doc.add_paragraph('看到以下输出表示成功：')
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

p = doc.add_paragraph()
run = p.add_run('''============================================================
🎉 部署成功！
============================================================

📱 访问地址：
   前端页面：http://你的服务器IP
   后端API：http://你的服务器IP/api

🗄️ 数据库信息：
   地址：localhost:3306
   数据库：ai_key_management
   用户名：aikey_user
   密码：（脚本自动生成的密码，请查看上方输出）

⚙️ 常用命令：
   查看日志：docker compose -f docker-compose.all-in-one.yml logs -f
   停止服务：docker compose -f docker-compose.all-in-one.yml down
   重启服务：docker compose -f docker-compose.all-in-one.yml restart

============================================================''')
set_chinese_font(run, 'Consolas', 9)

doc.add_page_break()

# 七、验证部署
doc.add_heading('七、验证部署', level=1)

doc.add_heading('7.1 浏览器访问', level=2)
p = doc.add_paragraph('打开浏览器，输入：')
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

p = doc.add_paragraph()
run = p.add_run('http://你的服务器IP')
set_chinese_font(run, 'Consolas', 10)

p = doc.add_paragraph('例如：http://123.45.67.89')
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

p = doc.add_paragraph('应该能看到登录页面')
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

doc.add_heading('7.2 测试登录', level=2)
p = doc.add_paragraph('默认账号：')
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

accounts = [
    '管理员：admin / admin123',
    '测试用户：test / test123'
]
for account in accounts:
    p = doc.add_paragraph(account, style='List Bullet')
    for run in p.runs:
        set_chinese_font(run, 'SimSun', 10)

doc.add_heading('7.3 检查服务状态', level=2)
p = doc.add_paragraph('在 FinalShell 执行：')
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

p = doc.add_paragraph()
run = p.add_run('cd /opt/aikey/deploy\ndocker compose -f docker-compose.all-in-one.yml ps')
set_chinese_font(run, 'Consolas', 10)

p = doc.add_paragraph('应该看到4个服务都是 Up 状态：')
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

p = doc.add_paragraph()
run = p.add_run('''NAME                STATUS
aikey-frontend      Up 5 minutes
aikey-backend       Up 5 minutes
aikey-mysql         Up 5 minutes
aikey-redis         Up 5 minutes''')
set_chinese_font(run, 'Consolas', 9)

doc.add_page_break()

# 八、常见问题
doc.add_heading('八、常见问题', level=1)

# Q1
doc.add_heading('Q1: 部署脚本运行失败', level=2)
p = doc.add_paragraph()
run = p.add_run('现象：')
set_chinese_font(run, 'SimSun', 10, True)
p.add_run('脚本报错退出')
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

p = doc.add_paragraph()
run = p.add_run('解决：')
set_chinese_font(run, 'SimSun', 10, True)
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

p = doc.add_paragraph()
run = p.add_run('''# 1. 检查是否在 deploy 目录
cd /opt/aikey/deploy
pwd  # 应该显示 /opt/aikey/deploy

# 2. 检查文件是否存在
ls -la finalshell-deploy.sh

# 3. 给脚本执行权限
chmod +x finalshell-deploy.sh

# 4. 重新运行
bash finalshell-deploy.sh''')
set_chinese_font(run, 'Consolas', 9)

# Q2
doc.add_heading('Q2: Docker安装失败', level=2)
p = doc.add_paragraph()
run = p.add_run('现象：')
set_chinese_font(run, 'SimSun', 10, True)
p.add_run('提示 "Failed to install Docker"')
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

p = doc.add_paragraph()
run = p.add_run('解决：')
set_chinese_font(run, 'SimSun', 10, True)
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

p = doc.add_paragraph()
run = p.add_run('''# 手动安装Docker
curl -fsSL https://get.docker.com | sh
systemctl start docker
systemctl enable docker
docker --version  # 验证安装

# 然后重新运行部署脚本
bash finalshell-deploy.sh''')
set_chinese_font(run, 'Consolas', 9)

# Q3
doc.add_heading('Q3: 端口被占用', level=2)
p = doc.add_paragraph()
run = p.add_run('现象：')
set_chinese_font(run, 'SimSun', 10, True)
p.add_run('提示 "port 80 is already allocated"')
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

p = doc.add_paragraph()
run = p.add_run('解决：')
set_chinese_font(run, 'SimSun', 10, True)
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

p = doc.add_paragraph()
run = p.add_run('''# 查看占用80端口的进程
netstat -tlnp | grep 80

# 停止占用端口的进程（如nginx、apache）
systemctl stop nginx
systemctl stop httpd

# 或者修改docker-compose使用其他端口
# 编辑 docker-compose.all-in-one.yml
# 将 "80:80" 改为 "8080:80"''')
set_chinese_font(run, 'Consolas', 9)

# Q4
doc.add_heading('Q4: 数据库连接失败', level=2)
p = doc.add_paragraph()
run = p.add_run('现象：')
set_chinese_font(run, 'SimSun', 10, True)
p.add_run('后端日志显示数据库连接错误')
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

p = doc.add_paragraph()
run = p.add_run('解决：')
set_chinese_font(run, 'SimSun', 10, True)
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

p = doc.add_paragraph()
run = p.add_run('''# 查看后端日志
docker compose -f docker-compose.all-in-one.yml logs backend

# 等待数据库完全启动（首次启动需要30-60秒）
sleep 60

# 重启后端服务
docker compose -f docker-compose.all-in-one.yml restart backend''')
set_chinese_font(run, 'Consolas', 9)

# Q5
doc.add_heading('Q5: 页面显示404', level=2)
p = doc.add_paragraph()
run = p.add_run('现象：')
set_chinese_font(run, 'SimSun', 10, True)
p.add_run('能访问但显示404错误')
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

p = doc.add_paragraph()
run = p.add_run('解决：')
set_chinese_font(run, 'SimSun', 10, True)
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

p = doc.add_paragraph()
run = p.add_run('''# 检查前端容器状态
docker compose -f docker-compose.all-in-one.yml ps frontend

# 查看前端日志
docker compose -f docker-compose.all-in-one.yml logs frontend

# 重启前端
docker compose -f docker-compose.all-in-one.yml restart frontend''')
set_chinese_font(run, 'Consolas', 9)

# Q6
doc.add_heading('Q6: 无法访问网站', level=2)
p = doc.add_paragraph()
run = p.add_run('现象：')
set_chinese_font(run, 'SimSun', 10, True)
p.add_run('浏览器打不开页面')
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

p = doc.add_paragraph()
run = p.add_run('排查步骤：')
set_chinese_font(run, 'SimSun', 10, True)
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

steps = [
    '检查安全组是否开放80端口',
    '检查服务器防火墙：\n   systemctl status firewalld\n   # 如果开启，添加规则：\n   firewall-cmd --permanent --add-port=80/tcp\n   firewall-cmd --reload',
    '检查服务状态：\n   docker compose -f docker-compose.all-in-one.yml ps'
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph(f'{i}. {step}')
    for run in p.runs:
        set_chinese_font(run, 'SimSun', 10)

doc.add_page_break()

# 九、日常维护
doc.add_heading('九、日常维护', level=1)

doc.add_heading('9.1 常用命令', level=2)
p = doc.add_paragraph()
run = p.add_run('''# 进入项目目录
cd /opt/aikey/deploy

# 查看所有服务状态
docker compose -f docker-compose.all-in-one.yml ps

# 查看日志（实时）
docker compose -f docker-compose.all-in-one.yml logs -f

# 查看特定服务日志
docker compose -f docker-compose.all-in-one.yml logs -f backend
docker compose -f docker-compose.all-in-one.yml logs -f frontend

# 重启所有服务
docker compose -f docker-compose.all-in-one.yml restart

# 重启特定服务
docker compose -f docker-compose.all-in-one.yml restart backend

# 停止所有服务
docker compose -f docker-compose.all-in-one.yml down

# 停止并删除数据（谨慎！）
docker compose -f docker-compose.all-in-one.yml down -v''')
set_chinese_font(run, 'Consolas', 9)

doc.add_heading('9.2 备份数据', level=2)
p = doc.add_paragraph()
run = p.add_run('''# 备份数据库
docker exec aikey-mysql mysqldump -u root -p ai_key_management > backup_$(date +%Y%m%d).sql

# 输入密码后，备份文件会在当前目录生成''')
set_chinese_font(run, 'Consolas', 9)

doc.add_heading('9.3 更新代码后重新部署', level=2)
p = doc.add_paragraph()
run = p.add_run('''cd /opt/aikey/deploy

# 1. 停止服务
docker compose -f docker-compose.all-in-one.yml down

# 2. 重新构建（代码有更新时）
docker compose -f docker-compose.all-in-one.yml up -d --build

# 3. 查看状态
docker compose -f docker-compose.all-in-one.yml ps''')
set_chinese_font(run, 'Consolas', 9)

doc.add_heading('9.4 查看资源使用', level=2)
p = doc.add_paragraph()
run = p.add_run('''# 查看容器资源使用
docker stats

# 查看服务器整体资源
df -h  # 磁盘
free -h  # 内存
top  # CPU和进程''')
set_chinese_font(run, 'Consolas', 9)

doc.add_page_break()

# 十、获取帮助
doc.add_heading('十、获取帮助', level=1)

p = doc.add_paragraph('如果遇到问题：')
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

steps = [
    '先查看日志：docker compose -f docker-compose.all-in-one.yml logs',
    '检查本文档的【常见问题】章节',
    '联系项目开发者'
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph(f'{i}. {step}')
    for run in p.runs:
        set_chinese_font(run, 'SimSun', 10)

doc.add_page_break()

# 附录：文件清单
doc.add_heading('附录：文件清单', level=1)

p = doc.add_paragraph('部署涉及的所有文件：')
for run in p.runs:
    set_chinese_font(run, 'SimSun', 10)

p = doc.add_paragraph()
run = p.add_run('''aikey/
├── backend/                          # 后端源代码
│   ├── src/
│   │   └── main/
│   │       ├── java/
│   │       └── resources/
│   └── pom.xml
├── frontend/                         # 前端源代码
│   ├── src/
│   ├── public/
│   ├── package.json
│   └── vite.config.js
└── deploy/                           # 部署配置
    ├── docker-compose.all-in-one.yml # 容器编排配置
    ├── Dockerfile.backend            # 后端构建配置
    ├── Dockerfile.frontend           # 前端构建配置
    ├── nginx.conf                    # Nginx反向代理配置
    ├── finalshell-deploy.sh          # 一键部署脚本
    ├── .env                          # 环境变量（自动生成）
    ├── .env.example                  # 环境变量模板
    └── 部署指南.md                    # 本文件''')
set_chinese_font(run, 'Consolas', 9)

# 文档信息
doc.add_paragraph('─' * 50)
p = doc.add_paragraph()
run = p.add_run('''文档版本：v1.0
最后更新：2026-04-07
适用项目：AI调度中心——企业API Key管理系统''')
set_chinese_font(run, 'SimSun', 10)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 保存文档
output_path = r'C:\Users\26404\Desktop\26计设材料\AI调度中心部署指南.docx'
doc.save(output_path)
print(f'文档已保存到: {output_path}')
